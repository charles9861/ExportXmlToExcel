#!/usr/bin/env python3
"""
python export_xml_to_word_v1.py
Parses clash XML and writes a formatted Word (.docx) document:
- Page setup: A3 landscape
- 8 columns with specific widths and styles
- Embeds images resized to fit within cells
Config: config.XML_FILE and config.OUTPUT_FILE
"""

import config
import xml.etree.ElementTree as ET
from pathlib import Path
import tempfile
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage

# ---------- Layout constants ----------
HEADER_FILL = "6D9DA7"
HEADER_FONT_SIZE = 14.5
DATA_FONT_SIZE = 10.5
HEADER_FONT_COLOR = (255, 255, 255)  # white
COL_WIDTHS_CM = [1, 4.5, 4.5, 4.5, 7.5, 7.5, 4.5, 4.5]

# ---------- Helper functions ----------
def find_image_file(href_raw: str, xml_path: Path):
    """Try multiple candidate paths and return Path or None."""
    if not href_raw:
        return None
    href = href_raw.replace("\\", "/").strip()
    candidates = []
    p = Path(href)
    if p.is_absolute():
        candidates.append(p)
    xml_dir = xml_path.parent
    candidates.append(xml_dir / href)
    fname = Path(href).name
    candidates.append(xml_dir / fname)
    candidates.append(xml_dir / "ELV_files" / fname)
    for cand in candidates:
        try:
            cand_resolved = cand.resolve()
        except Exception:
            cand_resolved = cand
        if cand_resolved.exists():
            return Path(cand_resolved)
    return None


def resize_image_for_cell(orig_path: Path, max_width_cm: float, max_height_cm: float):
    """Resize image preserving aspect ratio to fit within max dimensions in cm."""
    try:
        img = PILImage.open(orig_path)
    except Exception as e:
        print(f"Could not open image {orig_path}: {e}")
        return None
    max_w_px = int(max_width_cm * 96 / 2.54)
    max_h_px = int(max_height_cm * 96 / 2.54)
    w, h = img.size
    ratio = min(max_w_px / w, max_h_px / h, 1.0)
    new_w, new_h = int(w * ratio), int(h * ratio)
    img = img.resize((new_w, new_h), PILImage.LANCZOS)
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = Path(tf.name)
    tf.close()
    img.save(tmp_path)
    return tmp_path


def get_item_details(clash_object):
    """Return multiline item details (Item Name, Network, Item Type, Pipe sizes)."""
    if clash_object is None:
        return ""
    tags = {}
    for st in clash_object.findall(".//smarttag"):
        name = st.findtext("name", "").strip()
        val = st.findtext("value", "").strip()
        if name:
            tags[name] = val
    lines = []
    if tags.get("Item Name"):
        lines.append(f"Item Name: {tags.get('Item Name')}")
    if tags.get("Civil3D General:Network name"):
        lines.append(f"Network: {tags.get('Civil3D General:Network name')}")
    part = tags.get("Civil3D General:Part Size Name", "")
    itype = tags.get("Item Type", "")
    type_line = " ".join(filter(None, [part, itype])).strip()
    if type_line:
        lines.append(f"Item Type: {type_line}")
    inner = tags.get("Civil3D General:Inner Diameter or Width", "")
    outer = tags.get("Civil3D General:Outer Diameter or Width", "")
    if inner or outer:
        lines.append(f"Pipe {inner} x {outer}".strip())
    return "\n".join(lines)


def get_item_name_short(item_details_text):
    """Return just the Item Name for the 'Between' line (if present)."""
    for line in item_details_text.splitlines():
        if line.startswith("Item Name:"):
            return line.replace("Item Name:", "").strip()
    first = item_details_text.splitlines()[0] if item_details_text else ""
    return first or "Unknown"


# ---------- Main export function ----------
def export_to_word(xml_file, output_file):
    xml_path = Path(xml_file)
    if not xml_path.exists():
        print(f"XML file not found: {xml_file}")
        return

    tree = ET.parse(xml_path)
    root = tree.getroot()

    # Create Word document
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42.0)  # A3 width
    section.page_height = Cm(29.7)  # A3 height
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # Create table
    clashes = root.findall(".//clashresult")
    table = doc.add_table(rows=1, cols=8)
    table.autofit = False
    for i, width in enumerate(COL_WIDTHS_CM):
        table.columns[i].width = Cm(width)

    # Header row
    headers = [
        "RFI No.",
        "Clash details",
        "Clash entity 1",
        "Clash entity 2",
        "Clash screenshot",
        "Solution",
        "Description of Solution",
        "TSL comment",
    ]
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(HEADER_FONT_SIZE)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = hdr_cells[i]._element.xpath(".//w:shd")[0] if hdr_cells[i]._element.xpath(".//w:shd") else None
        if shading_elm is None:
            hdr_cells[i]._element.get_or_add_shd().fill = HEADER_FILL
        else:
            shading_elm.set("w:fill", HEADER_FILL)

    temp_images = []

        # Populate table
    for i, clash in enumerate(clashes, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)

        # Fill other cells, for example:
        clash_details = get_item_details(clash)
        row_cells[1].text = clash_details

            # Save document AFTER all rows are added
    from pathlib import Path
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")

        # RFI No.
    row_cells[0].text = str(i)

        # Clash details
    clash_group = "Unknown Group"
    clash_guid = clash.get("guid")
    for test in root.findall(".//clashtest"):
            for cr in test.findall(".//clashresult"):
                if clash_guid and cr.get("guid") == clash_guid:
                    clash_group = test.get("name", "Unknown Group")
                    break
            if clash_group != "Unknown Group":
                break

    clash_name = clash.get("name", f"Clash{i}")
    distance = clash.get("distance", "N/A")
    pos = clash.find(".//pos3f")
    coords_text = ""
    if pos is not None:
            x_val = float(pos.get("x") or 0.0)
            y_val = float(pos.get("y") or 0.0)
            z_val = float(pos.get("z") or 0.0)
            coords_text = f"{x_val:.3f}m, {y_val:.3f}m, {z_val:.3f}m"

            # ---------- Run ----------

    if __name__ == "__main__":
            import config
            export_to_word(config.XML_FILE, config.OUTPUT_FILE)

