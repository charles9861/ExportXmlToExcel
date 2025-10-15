#!/usr/bin/env python3
"""
python export_xml_to_word_v3.py
Parses clash XML and writes a formatted Word (.docx) document:
- Page setup: A3 landscape
- 8 columns with specific widths and styles
- Embeds images resized to fit within cells (optional)
Config: config.XML_FILE and config.OUTPUT_FILE
"""

import xml.etree.ElementTree as ET
from pathlib import Path
import tempfile
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import config

# ---------- Layout constants ----------
HEADER_FILL = "6D9DA7"
HEADER_FONT_SIZE = 14.5
DATA_FONT_SIZE = 10.5
HEADER_FONT_COLOR = (255, 255, 255)  # white
COL_WIDTHS_CM = [1, 4.5, 4.5, 4.5, 7.5, 7.5, 4.5, 4.5]

# ---------- Helper functions ----------
def find_image_file(href_raw: str, xml_path: Path):
    if not href_raw:
        return None
    href = href_raw.replace("\\", "/").strip()
    candidates = []
    p = Path(href)
    if p.is_absolute():
        candidates.append(p)
    xml_dir = xml_path.parent
    candidates.append(xml_dir / href)
    fname = Path(href).name
    candidates.append(xml_dir / fname)
    candidates.append(xml_dir / "ELV_files" / fname)
    for cand in candidates:
        try:
            cand_resolved = cand.resolve()
        except Exception:
            cand_resolved = cand
        if cand_resolved.exists():
            return Path(cand_resolved)
    return None

def resize_image_for_cell(orig_path: Path, max_width_cm: float, max_height_cm: float):
    try:
        img = PILImage.open(orig_path)
    except Exception as e:
        print(f"Could not open image {orig_path}: {e}")
        return None
    max_w_px = int(max_width_cm * 96 / 2.54)
    max_h_px = int(max_height_cm * 96 / 2.54)
    w, h = img.size
    ratio = min(max_w_px / w, max_h_px / h, 1.0)
    new_w, new_h = int(w * ratio), int(h * ratio)
    img = img.resize((new_w, new_h), PILImage.LANCZOS)
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = Path(tf.name)
    tf.close()
    img.save(tmp_path)
    return tmp_path

def get_item_details(clash_object):
    if clash_object is None:
        return ""
    tags = {}
    for st in clash_object.findall(".//smarttag"):
        name = st.findtext("name", "").strip()
        val = st.findtext("value", "").strip()
        if name:
            tags[name] = val
    lines = []
    if tags.get("Item Name"):
        lines.append(f"Item Name: {tags.get('Item Name')}")
    if tags.get("Civil3D General:Network name"):
        lines.append(f"Network: {tags.get('Civil3D General:Network name')}")
    part = tags.get("Civil3D General:Part Size Name", "")
    itype = tags.get("Item Type", "")
    type_line = " ".join(filter(None, [part, itype])).strip()
    if type_line:
        lines.append(f"Item Type: {type_line}")
    inner = tags.get("Civil3D General:Inner Diameter or Width", "")
    outer = tags.get("Civil3D General:Outer Diameter or Width", "")
    if inner or outer:
        lines.append(f"Pipe {inner} x {outer}".strip())
    return "\n".join(lines)

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), "auto")
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

# ---------- Main export function ----------
def export_to_word(xml_file, output_file):
    xml_path = Path(xml_file)
    if not xml_path.exists():
        print(f"XML file not found: {xml_file}")
        return

    tree = ET.parse(xml_path)
    root = tree.getroot()

    # Create Word document
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(1.0)

    # Create table
    clashes = root.findall(".//clashresult")
    table = doc.add_table(rows=1, cols=8)
    table.autofit = False
    for i, width in enumerate(COL_WIDTHS_CM):
        table.columns[i].width = Cm(width)

    # Header row
    headers = [
        "RFI No.",
        "Clash details",
        "Clash entity 1",
        "Clash entity 2",
        "Clash screenshot",
        "Solution",
        "Description of Solution",
        "TSL comment",
    ]
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(HEADER_FONT_SIZE)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], HEADER_FILL)

    # Populate table
    for i, clash in enumerate(clashes, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = get_item_details(clash)
        # You can fill additional cells as needed
        # e.g., row_cells[2].text = ...

    # Save document
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")

# ---------- Run script ----------
if __name__ == "__main__":
    export_to_word(config.XML_FILE, config.OUTPUT_FILE)

