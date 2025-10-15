#!/usr/bin/env python3
"""
python export_xml_to_word_v4.py
Parses clash XML and writes a formatted Word (.docx) document:
- Page setup: A3 landscape
- 8 columns with specific widths
- First 5 columns populated like Excel export: RFI No, Clash Details, Item1, Item2, Clash Image
- Images resized to fit cells
"""

import xml.etree.ElementTree as ET
from pathlib import Path
import tempfile
import os
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage
import config

# ---------- Layout constants ----------
HEADER_FILL = "2C7676"
HEADER_FONT_SIZE = 12
DATA_FONT_SIZE = 10
HEADER_FONT_COLOR = (255, 255, 255)
COL_WIDTHS_CM = [1, 4.5, 4.5, 4.5, 8, 8, 4.5, 4.5]  # match Excel-ish

# Image cell max size (cm)
IMG_MAX_W_CM = 7.5
IMG_MAX_H_CM = 7.5

# ---------- Helper functions ----------
def find_image_file(href_raw: str, xml_path: Path):
    if not href_raw:
        return None
    href = href_raw.replace("\\", "/").strip()
    candidates = []
    p = Path(href)
    if p.is_absolute():
        candidates.append(p)
    xml_dir = xml_path.parent
    candidates.append(xml_dir / href)
    fname = Path(href).name
    candidates.append(xml_dir / fname)
    candidates.append(xml_dir / "ELV_files" / fname)
    for cand in candidates:
        try:
            cand_resolved = cand.resolve()
        except Exception:
            cand_resolved = cand
        if cand_resolved.exists():
            return Path(cand_resolved)
    return None

def resize_image_for_cell(orig_path: Path, max_width_cm: float, max_height_cm: float):
    try:
        img = PILImage.open(orig_path)
    except Exception as e:
        print(f"Could not open image {orig_path}: {e}")
        return None
    # Word uses EMU internally; here we resize to fit roughly in cm
    max_w_px = int(max_width_cm * 96 / 2.54)
    max_h_px = int(max_height_cm * 96 / 2.54)
    w, h = img.size
    ratio = min(max_w_px / w, max_h_px / h, 1.0)
    new_w, new_h = int(w * ratio), int(h * ratio)
    img = img.resize((new_w, new_h), PILImage.LANCZOS)
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    tmp_path = Path(tf.name)
    tf.close()
    img.save(tmp_path)
    return tmp_path

def get_item_details(clash_object):
    if clash_object is None:
        return ""
    tags = {}
    for st in clash_object.findall(".//smarttag"):
        name = st.findtext("name", "").strip()
        val = st.findtext("value", "").strip()
        if name:
            tags[name] = val
    lines = []
    if tags.get("Item Name"):
        lines.append(f"Item Name: {tags.get('Item Name')}")
    if tags.get("Civil3D General:Network name"):
        lines.append(f"Network: {tags.get('Civil3D General:Network name')}")
    part = tags.get("Civil3D General:Part Size Name", "")
    itype = tags.get("Item Type", "")
    type_line = " ".join(filter(None, [part, itype])).strip()
    if type_line:
        lines.append(f"Item Type: {type_line}")
    inner = tags.get("Civil3D General:Inner Diameter or Width", "")
    outer = tags.get("Civil3D General:Outer Diameter or Width", "")
    if inner or outer:
        lines.append(f"Pipe {inner} x {outer}".strip())
    return "\n".join(lines)

def get_item_name_short(item_details_text):
    for line in item_details_text.splitlines():
        if line.startswith("Item Name:"):
            return line.replace("Item Name:", "").strip()
    first = item_details_text.splitlines()[0] if item_details_text else ""
    return first or "Unknown"

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), "auto")
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

# ---------- Main export ----------
def export_to_word(xml_file, output_file):
    xml_path = Path(xml_file)
    if not xml_path.exists():
        print(f"XML file not found: {xml_file}")
        return

    tree = ET.parse(xml_path)
    root = tree.getroot()

    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(42)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)

    clashes = root.findall(".//clashresult")
    table = doc.add_table(rows=1, cols=8)
    table.autofit = False
    for i, w in enumerate(COL_WIDTHS_CM):
        table.columns[i].width = Cm(w)

    # Headers
    headers = ["RFI No.", "Clash Details", "Item 1", "Item 2", "Clash Image", "Solution", "Description of Solution", "TSL comment"]
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(HEADER_FONT_SIZE)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*HEADER_FONT_COLOR)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(hdr_cells[i], HEADER_FILL)

    temp_images = []

    for i, clash in enumerate(clashes, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)

        # Clash group
        clash_group = "Unknown Group"
        clash_guid = clash.get("guid")
        for test in root.findall(".//clashtest"):
            for cr in test.findall(".//clashresult"):
                if clash_guid and cr.get("guid") == clash_guid:
                    clash_group = test.get("name", "Unknown Group")
                    break
            if clash_group != "Unknown Group":
                break

        # Clash basic
        clash_name = clash.get("name", f"Clash{i}")
        distance = clash.get("distance", "N/A")
        pos = clash.find(".//pos3f")
        coords_text = ""
        if pos is not None:
            x_val = float(pos.get("x") or 0)
            y_val = float(pos.get("y") or 0)
            z_val = float(pos.get("z") or 0)
            coords_text = f"{x_val:.3f}m, {y_val:.3f}m, {z_val:.3f}m"

        # Items
        objs = clash.findall(".//clashobject")
        item1_text = get_item_details(objs[0]) if len(objs) > 0 else ""
        item2_text = get_item_details(objs[1]) if len(objs) > 1 else ""
        between_line = f"Between: {get_item_name_short(item1_text)} and {get_item_name_short(item2_text)}\n"

        clash_details = f"Clash Group: {clash_group}\n{between_line}{clash_name}\nDistance: {distance}m\nClash Point: {coords_text}"
        row_cells[1].text = clash_details
        row_cells[2].text = item1_text
        row_cells[3].text = item2_text

        # Clash image
        href_raw = clash.get("href") or ""
        img_path = find_image_file(href_raw, xml_path)
        if img_path:
            tmp_img = resize_image_for_cell(img_path, IMG_MAX_W_CM, IMG_MAX_H_CM)
            if tmp_img:
                run = row_cells[4].paragraphs[0].add_run()
                run.add_picture(str(tmp_img), width=Cm(IMG_MAX_W_CM))
                temp_images.append(tmp_img)
            else:
                row_cells[4].text = str(img_path)
        else:
            row_cells[4].text = href_raw or ""

    # Ensure output folder exists
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")

    # cleanup temp images
    for t in temp_images:
        try:
            os.remove(t)
        except Exception:
            pass

# ---------- Run ----------
if __name__ == "__main__":
    export_to_word(config.XML_FILE, config.OUTPUT_FILE)
